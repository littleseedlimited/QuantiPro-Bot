from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from enum import Enum
import os
from .citations import CitationManager, Reference, CitationStyle


class DocumentStructure(Enum):
    """Document structure templates."""
    IMRAD = "imrad"           # Introduction, Methods, Results, And Discussion
    APA_RESEARCH = "apa"      # APA style research paper
    THESIS = "thesis"         # Thesis/Dissertation format
    REPORT = "report"         # General research report
    JOURNAL = "journal"       # Journal article format
    CUSTOM = "custom"         # User-defined sections


class FontFamily(Enum):
    """Supported font families."""
    TIMES_NEW_ROMAN = "Times New Roman"
    ARIAL = "Arial"
    CALIBRI = "Calibri"
    GEORGIA = "Georgia"
    CAMBRIA = "Cambria"
    GARAMOND = "Garamond"
    COURIER_NEW = "Courier New"


class LineSpacing(Enum):
    """Line spacing options."""
    SINGLE = 1.0
    ONE_HALF = 1.5
    DOUBLE = 2.0
    CUSTOM = 0  # Use custom value


@dataclass
class ManuscriptSettings:
    """
    Comprehensive manuscript formatting settings.
    """
    # Font Settings
    font_family: FontFamily = FontFamily.TIMES_NEW_ROMAN
    font_size: int = 12  # in points
    heading_font_size: int = 14
    title_font_size: int = 16
    
    # Spacing
    line_spacing: LineSpacing = LineSpacing.DOUBLE
    custom_line_spacing: float = 2.0
    paragraph_spacing_after: int = 12  # in points
    
    # Margins (in inches)
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.0
    margin_right: float = 1.0
    
    # Document Structure
    structure: DocumentStructure = DocumentStructure.IMRAD
    
    # Word Count
    target_word_count: Optional[int] = None
    max_word_count: Optional[int] = None
    
    # Citation Style
    citation_style: CitationStyle = CitationStyle.APA7
    
    # Section Headers (for IMRAD)
    include_abstract: bool = True
    include_keywords: bool = False
    include_acknowledgments: bool = False
    include_appendix: bool = False
    
    # Page Numbers
    include_page_numbers: bool = True
    page_number_position: str = "bottom-center"  # bottom-center, bottom-right, top-right
    
    # Additional Options
    justify_text: bool = False
    first_line_indent: float = 0.5  # in inches


# Structure templates defining section order
STRUCTURE_TEMPLATES = {
    DocumentStructure.IMRAD: [
        "Title Page",
        "Abstract", 
        "Introduction",
        "Methods",
        "Results",
        "Discussion",
        "Conclusion",
        "References"
    ],
    DocumentStructure.APA_RESEARCH: [
        "Title Page",
        "Abstract",
        "Introduction",
        "Literature Review",
        "Methods",
        "Results",
        "Discussion",
        "References",
        "Appendix"
    ],
    DocumentStructure.THESIS: [
        "Title Page",
        "Declaration",
        "Abstract",
        "Acknowledgments",
        "Table of Contents",
        "List of Tables",
        "List of Figures",
        "Chapter 1: Introduction",
        "Chapter 2: Literature Review",
        "Chapter 3: Methodology",
        "Chapter 4: Results",
        "Chapter 5: Discussion",
        "Chapter 6: Conclusion",
        "References",
        "Appendices"
    ],
    DocumentStructure.REPORT: [
        "Title Page",
        "Executive Summary",
        "Introduction",
        "Background",
        "Analysis",
        "Findings",
        "Recommendations",
        "Conclusion",
        "References"
    ],
    DocumentStructure.JOURNAL: [
        "Title",
        "Authors",
        "Abstract",
        "Keywords",
        "Introduction",
        "Materials and Methods",
        "Results",
        "Discussion",
        "Conclusion",
        "Acknowledgments",
        "References"
    ]
}


class ManuscriptGenerator:
    """
    Advanced manuscript generator with comprehensive formatting options.
    """
    
    def __init__(self, settings: ManuscriptSettings = None):
        self.settings = settings or ManuscriptSettings()
        self.doc = Document()
        self.word_count = 0
        self._setup_document()
    
    def _setup_document(self):
        """Configure document with settings."""
        # Set margins
        for section in self.doc.sections:
            section.top_margin = Inches(self.settings.margin_top)
            section.bottom_margin = Inches(self.settings.margin_bottom)
            section.left_margin = Inches(self.settings.margin_left)
            section.right_margin = Inches(self.settings.margin_right)
        
        # Configure Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.settings.font_family.value
        font.size = Pt(self.settings.font_size)
        
        # Line spacing
        para_format = style.paragraph_format
        if self.settings.line_spacing == LineSpacing.DOUBLE:
            para_format.line_spacing = 2.0
        elif self.settings.line_spacing == LineSpacing.ONE_HALF:
            para_format.line_spacing = 1.5
        elif self.settings.line_spacing == LineSpacing.SINGLE:
            para_format.line_spacing = 1.0
        else:
            para_format.line_spacing = self.settings.custom_line_spacing
        
        para_format.space_after = Pt(self.settings.paragraph_spacing_after)
        
        # Configure heading styles
        for level in range(1, 4):
            try:
                heading_style = self.doc.styles[f'Heading {level}']
                heading_style.font.name = self.settings.font_family.value
                heading_style.font.size = Pt(self.settings.heading_font_size - (level - 1) * 2)
                heading_style.font.bold = True
            except:
                pass
    
    def _add_paragraph(self, text: str, bold: bool = False, italic: bool = False, 
                       center: bool = False, font_size: int = None) -> None:
        """Add a paragraph with formatting and track word count."""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if font_size:
            run.font.size = Pt(font_size)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif self.settings.justify_text:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # First line indent
        if self.settings.first_line_indent > 0 and not center:
            p.paragraph_format.first_line_indent = Inches(self.settings.first_line_indent)
        
        # Track word count
        self.word_count += len(text.split())
        
        return p
    
    
    def _add_bullet_point(self, text: str) -> None:
        """Add a paragraph with bullet point style."""
        p = self.doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        
        # Track word count
        self.word_count += len(text.split())
        return p
    
    def _add_table(self, data: Any, title: str = "") -> None:
        """Render a pandas DataFrame or list of dicts as a Word table."""
        import pandas as pd
        
        if title:
            self._add_paragraph(title, bold=True, italic=True, center=True)
            
        df = None
        if isinstance(data, pd.DataFrame):
            df = data.reset_index() if data.index.name else data
        elif isinstance(data, dict):
            df = pd.DataFrame(data)
        elif isinstance(data, list):
            df = pd.DataFrame(data)
            
        if df is not None:
            # Create table
            table = self.doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'
            
            # Header
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = str(col)
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                
            # Rows
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    val_str = str(value)
                    if isinstance(value, float):
                        val_str = f"{value:.3f}"
                    row_cells[i].text = val_str
            
            self.doc.add_paragraph()  # Spacing after table
    
    def generate(self, 
                 filename: str,
                 title: str,
                 authors: List[str],
                 abstract: str,
                 content_sections: Dict[str, str] = None,
                 stats_results: List[str] = None,
                 discussion_text: str = None,
                 references: List[Reference] = None,
                 images: List[str] = None,
                 keywords: List[str] = None,
                 acknowledgments: str = None,
                 methods_text: str = None,
                 conclusion_text: str = None):
        """
        Generate a complete manuscript with comprehensive formatting.
        
        Args:
            filename: Output file path
            title: Manuscript title
            authors: List of author names
            abstract: Abstract text
            content_sections: Additional content sections
            stats_results: Statistical results
            discussion_text: AI-generated discussion
            references: Bibliography references
            images: Figure image paths
            keywords: Research keywords
            acknowledgments: Acknowledgments text
            methods_text: Methods section content
            conclusion_text: Conclusion section content
        """
        content_sections = content_sections or {}
        
        # Title Page
        self._add_title_page(title, authors)
        self.doc.add_page_break()
        
        # Abstract (if enabled)
        if self.settings.include_abstract and abstract:
            self._add_abstract(abstract, keywords)
            self.doc.add_page_break()
        
        # Main Content based on structure
        if self.settings.structure == DocumentStructure.IMRAD:
            self._generate_imrad(content_sections, methods_text, stats_results, 
                                discussion_text, conclusion_text)
        else:
            # Custom/General structure
            for header, content in content_sections.items():
                self.doc.add_heading(header, level=1)
                for para in content.split('\n\n'):
                    if para.strip():
                        self._add_paragraph(para.strip())
            
            if stats_results:
                self.doc.add_heading("Results", level=1)
                for res in stats_results:
                    for para in res.split('\n\n'):
                        if para.strip():
                            self._add_paragraph(para.strip())
            
            if discussion_text:
                self.doc.add_page_break()
                self.doc.add_heading("Discussion", level=1)
                for para in discussion_text.split('\n\n'):
                    if para.strip():
                        if para.strip().startswith('•') or para.strip().startswith('-'):
                             self._add_bullet_point(para.strip()[1:].strip())
                        else:
                             self._add_paragraph(para.strip())
        
        # Conclusion (if provided separately)
        if conclusion_text and self.settings.structure != DocumentStructure.IMRAD:
            self.doc.add_heading("Conclusion", level=1)
            for para in conclusion_text.split('\n\n'):
                if para.strip():
                    self._add_paragraph(para.strip())
        
        # Acknowledgments
        if self.settings.include_acknowledgments and acknowledgments:
            self.doc.add_heading("Acknowledgments", level=1)
            self._add_paragraph(acknowledgments)
        
        # Figures
        if images:
            self._add_figures(images)
        
        # References
        if references:
            self._add_references(references)
        
        # Add word count info
        self._add_word_count_footer()
        
        self.doc.save(filename)
        return filename, self.word_count
    
    
    def _generate_imrad(self, content_sections: Dict, methods_text: str,
                        stats_results: List[Any], discussion_text: str,
                        conclusion_text: str):
        """Generate IMRAD format document."""
        
        # Introduction
        intro = content_sections.get('Introduction') or content_sections.get('Research Objectives', '')
        if intro:
            self.doc.add_heading("Introduction", level=1)
            # Handle comma separated objectives/intro
            if ',' in intro and (content_sections.get('Research Objectives') == intro or 'Objectives' in intro):
                items = [i.strip() for i in intro.split(',') if i.strip()]
                self._add_paragraph("The primary objectives of this research include:")
                for item in items:
                    self._add_bullet_point(item)
            else:
                for para in intro.split('\n\n'):
                    if para.strip():
                        self._add_paragraph(para.strip())
        
        # Research Questions if provided in content_sections
        questions = content_sections.get('Research Questions')
        if questions:
             self.doc.add_heading("Research Questions", level=2)
             if ',' in questions:
                 items = [i.strip() for i in questions.split(',') if i.strip()]
                 for item in items:
                     self._add_bullet_point(item)
             else:
                 self._add_paragraph(questions)
        
        # Methods
        methods = methods_text or content_sections.get('Methods') or content_sections.get('Methodology', '')
        if methods:
            self.doc.add_heading("Methods", level=1)
            for para in methods.split('\n\n'):
                if para.strip():
                    self._add_paragraph(para.strip())
        
        # Results
        if stats_results:
            self.doc.add_heading("Results", level=1)
            for res in stats_results:
                if isinstance(res, dict) and 'type' in res and res['type'] == 'table':
                    # Structured table data
                    self._add_table(res['data'], title=res.get('title', ''))
                    if 'narrative' in res:
                         self._add_paragraph(res['narrative'])
                elif isinstance(res, str):
                    # Legacy string format
                    for para in res.split('\n\n'):
                        if para.strip():
                            self._add_paragraph(para.strip())
        
        # Discussion
        if discussion_text:
            self.doc.add_page_break()
            self.doc.add_heading("Discussion", level=1)
            for para in discussion_text.split('\n\n'):
                if para.strip():
                    self._add_paragraph(para.strip())
        
        # Conclusion
        conclusion = conclusion_text or content_sections.get('Conclusion', '')
        if conclusion:
            self.doc.add_heading("Conclusion", level=1)
            for para in conclusion.split('\n\n'):
                if para.strip():
                    self._add_paragraph(para.strip())
    
    def _add_title_page(self, title: str, authors: List[str]):
        """Create formatted title page."""
        # Spacing before title
        for _ in range(5):
            self.doc.add_paragraph()
        
        # Title
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(self.settings.title_font_size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Authors
        p = self.doc.add_paragraph()
        run = p.add_run(", ".join(authors))
        run.font.size = Pt(self.settings.font_size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        from datetime import datetime
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        run = p.add_run(datetime.now().strftime("%B %Y"))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_abstract(self, text: str, keywords: List[str] = None):
        """Create abstract section."""
        # Abstract header
        p = self.doc.add_paragraph()
        run = p.add_run("Abstract")
        run.bold = True
        run.font.size = Pt(self.settings.heading_font_size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()
        
        # Abstract text
        self._add_paragraph(text)
        
        # Keywords
        if self.settings.include_keywords and keywords:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run1 = p.add_run("Keywords: ")
            run1.bold = True
            run2 = p.add_run(", ".join(keywords))
            run2.italic = True
    
    def _add_figures(self, images: List[str]):
        """Add figures section."""
        self.doc.add_heading("Figures", level=1)
        
        for i, img_path in enumerate(images, 1):
            if img_path and os.path.exists(img_path):
                try:
                    self.doc.add_picture(img_path, width=Inches(6.0))
                    last_p = self.doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    caption = self.doc.add_paragraph()
                    run = caption.add_run(f"Figure {i}: {os.path.basename(img_path)}")
                    run.italic = True
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    self._add_paragraph(f"[Error embedding figure {i}: {e}]")
    
    def _add_references(self, references: List[Reference]):
        """Add formatted references section."""
        self.doc.add_page_break()
        self.doc.add_heading("References", level=1)
        
        for ref in references:
            formatted = CitationManager.format_entry(ref, self.settings.citation_style)
            p = self.doc.add_paragraph(formatted)
            # Hanging indent for references
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.left_indent = Inches(0.5)
    
    def _add_word_count_footer(self):
        """Add word count information."""
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        
        word_info = f"Word Count: {self.word_count}"
        
        if self.settings.target_word_count:
            diff = self.word_count - self.settings.target_word_count
            if diff > 0:
                word_info += f" (+{diff} over target of {self.settings.target_word_count})"
            elif diff < 0:
                word_info += f" ({abs(diff)} under target of {self.settings.target_word_count})"
            else:
                word_info += f" (target: {self.settings.target_word_count} ✓)"
        
        if self.settings.max_word_count and self.word_count > self.settings.max_word_count:
            word_info += f" ⚠️ EXCEEDS MAX ({self.settings.max_word_count})"
        
        run = p.add_run(word_info)
        run.italic = True
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def get_available_settings():
    """Return formatted string of available formatting options."""
    return {
        'fonts': [f.value for f in FontFamily],
        'structures': [s.value for s in DocumentStructure],
        'citation_styles': [c.value for c in CitationStyle],
        'line_spacing': ['Single', '1.5', 'Double'],
        'font_sizes': [10, 11, 12, 14],
    }
